from zipfile import ZIP_DEFLATED, ZipFile
from xml.etree import ElementTree

from docx import Document


def test_docx_메타데이터를_비우고_ooxml_무결성을_검증한다(tmp_path):
    from document_convert.docx_privacy import scrub_and_validate

    source = tmp_path / "source.docx"
    document = Document()
    document.add_paragraph("example.com fixture")
    document.core_properties.author = "Example Author"
    document.core_properties.title = "Private title"
    document.save(source)
    target = tmp_path / "clean.docx"

    scrub_and_validate(source, target)

    with ZipFile(target) as archive:
        assert archive.testzip() is None
    cleaned = Document(target).core_properties
    assert cleaned.author == ""
    assert cleaned.title == ""
    assert cleaned.last_modified_by == ""
    with ZipFile(target) as archive:
        app_properties = ElementTree.fromstring(archive.read("docProps/app.xml"))
    company = next(
        (element.text or "" for element in app_properties.iter() if element.tag.endswith("}Company")),
        "",
    )
    assert company == ""
    for name in ("Manager", "HyperlinkBase"):
        assert not any(element.tag.endswith("}" + name) and (element.text or "") for element in app_properties.iter())
    assert "docProps/custom.xml" not in archive.namelist()

    with ZipFile(target) as archive:
        relationships = archive.read("_rels/.rels")
        content_types = archive.read("[Content_Types].xml")
        assert b"/docProps/custom.xml" not in relationships
        assert b"/docProps/custom.xml" not in content_types


def test_손상된_docx는_검증실패한다(tmp_path):
    import pytest
    from document_convert.docx_privacy import InvalidDocxError, scrub_and_validate

    source = tmp_path / "broken.docx"
    source.write_bytes(b"not a zip")
    with pytest.raises(InvalidDocxError):
        scrub_and_validate(source, tmp_path / "clean.docx")


def test_custom_properties와_관계_콘텐츠타입을_실제로_제거한다(tmp_path):
    from document_convert.docx_privacy import scrub_and_validate

    source = tmp_path / "with-custom.docx"
    document = Document()
    document.add_paragraph("example.com fixture")
    document.save(source)
    custom_xml = b"<?xml version='1.0'?><Properties xmlns='http://schemas.openxmlformats.org/officeDocument/2006/custom-properties'><property name='FixtureSecret'/></Properties>"
    relationship = b"<Relationships xmlns='http://schemas.openxmlformats.org/package/2006/relationships'><Relationship Id='rId4' Type='http://schemas.openxmlformats.org/officeDocument/2006/relationships/custom-properties' Target='docProps/custom.xml'/></Relationships>"
    content_type = b"<Types xmlns='http://schemas.openxmlformats.org/package/2006/content-types'><Override PartName='/docProps/custom.xml' ContentType='application/vnd.openxmlformats-officedocument.custom-properties+xml'/></Types>"
    staged = tmp_path / "staged.docx"
    with ZipFile(source) as original, ZipFile(staged, "w", ZIP_DEFLATED) as updated:
        for item in original.infolist():
            if item.filename == "_rels/.rels":
                updated.writestr(item, relationship)
            elif item.filename == "[Content_Types].xml":
                updated.writestr(item, content_type)
            else:
                updated.writestr(item, original.read(item.filename))
        updated.writestr("docProps/custom.xml", custom_xml)
    target = tmp_path / "clean.docx"

    scrub_and_validate(staged, target)

    with ZipFile(target) as cleaned:
        assert cleaned.testzip() is None
        assert "docProps/custom.xml" not in cleaned.namelist()
        assert b"docProps/custom.xml" not in cleaned.read("_rels/.rels")
        assert b"docProps/custom.xml" not in cleaned.read("[Content_Types].xml")


def test_validate_docx는_예상텍스트를_검증한다(tmp_path):
    from document_convert.docx_privacy import validate_docx

    source = tmp_path / "valid.docx"
    document = Document()
    document.add_paragraph("Synthetic example.com text")
    document.save(source)

    assert validate_docx(source, expected_text=["Synthetic", "example.com"]) is True


def test_core_properties의_revision은_유효한_양의정수이고_식별필드와_날짜는_제거된다(tmp_path):
    from datetime import datetime, timezone

    from document_convert.docx_privacy import scrub_and_validate

    source = tmp_path / "metadata.docx"
    document = Document()
    document.add_paragraph("example.com fixture")
    properties = document.core_properties
    properties.author = "Example Author"
    properties.last_modified_by = "Example Editor"
    properties.revision = 7
    properties.created = datetime.now(timezone.utc)
    properties.modified = datetime.now(timezone.utc)
    document.save(source)
    target = tmp_path / "clean-metadata.docx"

    scrub_and_validate(source, target)

    with ZipFile(target) as archive:
        core = ElementTree.fromstring(archive.read("docProps/core.xml"))
    values = {element.tag.rsplit("}", 1)[-1]: (element.text or "") for element in core}
    assert values.get("revision", "").isdigit()
    assert int(values["revision"]) > 0
    assert values.get("creator", "") == ""
    assert values.get("lastModifiedBy", "") == ""
    assert values.get("created", "") == ""
    assert values.get("modified", "") == ""
